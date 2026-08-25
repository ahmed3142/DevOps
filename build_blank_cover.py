#!/usr/bin/env python3
"""Builds a blank, reusable Assignment Cover Sheet in Word format.

Reuses the layout helpers from build_docx.py so the blank template and the
one bound into the report cannot drift apart.
Run: ~/.ctf-tools/venv/bin/python3 build_blank_cover.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

import build_docx as bd

OUT = Path(__file__).parent / "out" / "Assignment_Cover_Sheet_BLANK.docx"

FIELDS = [
    ("Learner Name and Surname:", "", False),
    ("Learner Registration No.", "", False),
    ("Study Centre Name", "Learn Key Institute", True),
    ("Qualification Title", "", False),
    ("Unit Reference No.", "", False),
    ("Unit Title", "", False),
    ("Submission Date", "", False),
]


def main() -> int:
    doc = Document()
    bd.base_styles(doc)
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.25)
    s.top_margin = s.bottom_margin = Cm(2.0)

    t = bd.plain(doc, "Assignment Cover Sheet", size=16, bold=True, italic=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    t.runs[0].font.name = "Times New Roman"

    logo = bd.ROOT / "logos" / "learnkey.png"
    if logo.exists():
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.paragraph_format.space_after = Pt(12)
        lp.add_run().add_picture(str(logo), height=Cm(1.75))

    i = bd.plain(doc, "This cover sheet must be completed and added to the front of every assignment",
                 size=11, bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
    i.runs[0].font.name = "Times New Roman"

    tbl = doc.add_table(rows=0, cols=2)
    for label, value, bold_value in FIELDS:
        row = tbl.add_row()
        bd.cell_text(row.cells[0], label, bold=True, italic=True, space=4)
        bd.cell_text(row.cells[1], value, bold=bold_value, space=4)
        bd.box_borders(row.cells[0])
        bd.box_borders(row.cells[1])
    tbl.columns[0].width = Cm(6.2)
    tbl.columns[1].width = Cm(10.3)

    dec = doc.add_table(rows=1, cols=1).rows[0].cells[0]
    dec.text = ""
    head = dec.paragraphs[0]
    head.paragraph_format.space_after = Pt(4)
    hr = head.add_run("Declaration of authenticity:")
    hr.bold = hr.italic = True
    hr.font.size = Pt(10.5)
    hr.font.name = "Times New Roman"
    for n, line in enumerate([
        "I declare that the attached submission is my own original work. No significant part of it has been "
        "submitted for any other assignment and I have acknowledged in my notes and bibliography all written "
        "and electronic sources used.",
        "I acknowledge that my assignment will be subject to electronic scrutiny for academic honesty.",
        "I understand that failure to meet these guidelines may instigate the centre's malpractice procedures "
        "and risk failure of the unit and/or qualification.",
    ], start=1):
        par = dec.add_paragraph()
        par.paragraph_format.space_after = Pt(4)
        par.paragraph_format.left_indent = Cm(0.6)
        r = par.add_run(f"{n}.  {line}")
        r.bold = r.italic = True
        r.font.size = Pt(10.5)
        r.font.name = "Times New Roman"
    bd.box_borders(dec)

    sig = doc.add_table(rows=1, cols=2).rows[0]
    for idx, label in enumerate(("Learner signature", "Tutor signature")):
        cell = sig.cells[idx]
        cell.text = ""
        cell.paragraphs[0].paragraph_format.space_after = Pt(20)
        for text, is_rule in (("__________________________", True), (label, False), ("Date:", False)):
            par = cell.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.space_after = Pt(1)
            r = par.add_run(text)
            r.font.size = Pt(10.5)
            r.font.name = "Times New Roman"
            r.bold = r.italic = not is_rule
        bd.box_borders(cell)

    note = doc.add_table(rows=1, cols=1).rows[0].cells[0]
    bd.cell_text(note, "Note: Assignments must be submitted in typed PDF format only; handwritten "
                       "assignments are not accepted.", size=10)
    bd.box_borders(note)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"built: out/{OUT.name} ({OUT.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
