#!/usr/bin/env python3
"""Builds the Word version of the report from the same HTML sources.

Run through the virtualenv that has python-docx:
    ~/.ctf-tools/venv/bin/python3 build_docx.py
"""
from __future__ import annotations

import re
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).parent
OUT = ROOT / "out" / "DevOps_Delivery_Strategy.docx"


def _assessed_word_count() -> str:
    """Reads the count stamped onto report.html by assemble.py."""
    import re
    m = re.search(r"([\d,]+) words \(excluding", (ROOT / "report.html").read_text(encoding="utf-8"))
    return m.group(1) if m else "—"


WORD_COUNT = _assessed_word_count()

BODY_FONT = "Cambria"
HEAD_FONT = "Cambria"
MONO_FONT = "Consolas"
INK = RGBColor(0x14, 0x16, 0x1A)
MUTED = RGBColor(0x5B, 0x62, 0x6B)
ACCENT = RGBColor(0x7B, 0x2D, 0x26)
RULE = "D6D9DD"
DARK = "14161A"

# Figures are inline SVG in the HTML; the DOCX uses the rasterised copies.
FIGURE_IMAGES = {
    "fig1-value-stream": ("shots/fig1-value-stream.png", Cm(16.5)),
    "fig2-pipeline": ("shots/fig2-pipeline.png", Cm(16.5)),
    "fig3-ci-run": ("shots/ci-run.png", Cm(16.5)),
    "fig3-architecture": ("shots/fig3-architecture.png", Cm(16.5)),
    "fig4-error-budget": ("shots/fig4-error-budget.png", Cm(16.5)),
    "fig6-app-ui": ("shots/app-ui.png", Cm(9.2)),
}


# --------------------------------------------------------------------------
# HTML -> block list
# --------------------------------------------------------------------------
class Block:
    def __init__(self, kind, **kw):
        self.kind = kind
        self.__dict__.update(kw)


class Parser(HTMLParser):
    """Turns the report HTML into a flat list of blocks the writer understands."""

    VOID = {"br", "img", "meta", "link", "hr", "input"}
    SKIP_SUBTREE = {"svg", "style", "script"}
    INLINE = {"strong", "b", "em", "i", "code", "a", "span"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocks: list[Block] = []
        self.stack: list[str] = []
        self.runs: list[tuple[str, set]] = []
        self.fmt: set[str] = set()
        self.capture = None
        self.table = None
        self.row = None
        self.cell = None
        self.list_items = None
        self.note = None
        self.skip_depth = 0
        self.svg_depth = 0

    # -- helpers -------------------------------------------------------
    def _flush(self, kind, **kw):
        runs = [(t, set(f)) for t, f in self.runs if t.strip() or t == " "]
        self.runs = []
        if not runs and kind not in {"figure"}:
            return
        target = self.note["body"] if self.note is not None else self.blocks
        target.append(Block(kind, runs=runs, **kw))

    def handle_starttag(self, tag, attrs):
        a = dict(attrs)
        classes = a.get("class", "").split()

        if self.svg_depth or tag in self.SKIP_SUBTREE:
            if tag not in self.VOID:
                self.svg_depth += 1
            return

        if tag in self.VOID:
            if tag == "img" and self.capture == "figure":
                self.fig_src = a.get("src", "")
            return

        if self.skip_depth or "no-count" in classes and tag == "div" and "cover" in classes:
            self.skip_depth += 1
            self.stack.append(tag)
            return

        self.stack.append(tag)

        if tag in self.INLINE:
            if tag in {"strong", "b"}:
                self.fmt.add("b")
            elif tag in {"em", "i"}:
                self.fmt.add("i")
            elif tag == "code":
                self.fmt.add("code")
            return

        if tag == "table":
            self.table = {"caption": "", "head": [], "rows": []}
        elif tag == "tr":
            self.row = []
        elif tag in {"td", "th"}:
            self.cell = {"text": "", "num": "num" in classes, "head": tag == "th"}
        elif tag == "caption":
            self.capture = "caption"
        elif tag in {"ul", "ol"} and "refs" not in classes:
            self.list_items = []
        elif tag == "li":
            self.runs = []
        elif tag == "figure":
            self.capture = "figure"
            self.fig_src = ""
        elif tag == "figcaption":
            self.runs = []
        elif tag == "div" and "note" in classes:
            self.note = {"body": []}
        elif tag == "pre":
            self.capture = "pre"
            self.pre_text = ""
        elif tag == "section":
            self.runs = []

    def handle_endtag(self, tag):
        if tag in self.VOID:
            return
        if self.svg_depth:
            self.svg_depth -= 1
            return
        if self.skip_depth:
            if self.stack:
                self.stack.pop()
            self.skip_depth -= 1
            return
        if self.stack:
            self.stack.pop()

        if tag in {"strong", "b"}:
            self.fmt.discard("b")
        elif tag in {"em", "i"}:
            self.fmt.discard("i")
        elif tag == "code":
            self.fmt.discard("code")
        elif tag in {"h2", "h3", "h4"}:
            self._flush(tag)
        elif tag == "p":
            style = "lead" if "lead" in (self.current_class or "") else "p"
            self._flush(style)
        elif tag == "span":
            pass
        elif tag == "caption":
            self.capture = None
            self.table["caption"] = "".join(t for t, _ in self.runs).strip()
            self.runs = []
        elif tag in {"td", "th"}:
            self.cell["text"] = "".join(t for t, _ in self.runs).strip()
            self.runs = []
            self.row.append(self.cell)
            self.cell = None
        elif tag == "tr":
            if self.row and self.row[0]["head"]:
                self.table["head"] = self.row
            else:
                self.table["rows"].append(self.row)
            self.row = None
        elif tag == "table":
            target = self.note["body"] if self.note is not None else self.blocks
            target.append(Block("table", **self.table))
            self.table = None
        elif tag == "li":
            if self.list_items is not None:
                self.list_items.append([(t, set(f)) for t, f in self.runs])
            self.runs = []
        elif tag in {"ul", "ol"}:
            if self.list_items:
                target = self.note["body"] if self.note is not None else self.blocks
                target.append(Block("list", items=self.list_items))
            self.list_items = None
        elif tag == "figcaption":
            self.fig_caption = [(t, set(f)) for t, f in self.runs]
            self.runs = []
        elif tag == "figure":
            self.capture = None
            self.blocks.append(
                Block("figure", src=self.fig_src, caption=getattr(self, "fig_caption", []))
            )
            self.fig_caption = []
        elif tag == "pre":
            self.capture = None
            self.blocks.append(Block("pre", text=self.pre_text.strip("\n")))
        elif tag == "div" and self.note is not None:
            self.blocks.append(Block("note", body=self.note["body"]))
            self.note = None

    current_class = None

    def handle_data(self, data):
        if self.skip_depth or self.svg_depth:
            return
        if self.capture == "pre":
            self.pre_text += data
            return
        if not data.strip():
            if self.runs:
                self.runs.append((" ", set(self.fmt)))
            return
        text = re.sub(r"\s+", " ", data)
        self.runs.append((text, set(self.fmt)))


# The parser needs to know the class of the <p> it is closing; track it simply.
_orig_start = Parser.handle_starttag


def _start_with_class(self, tag, attrs):
    if tag == "p":
        self.current_class = dict(attrs).get("class", "")
    _orig_start(self, tag, attrs)


Parser.handle_starttag = _start_with_class


# --------------------------------------------------------------------------
# Word writing helpers
# --------------------------------------------------------------------------
def shade(cell_or_par, hex_colour):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_colour)
    cell_or_par.get_or_add_tcPr().append(el) if hasattr(cell_or_par, "get_or_add_tcPr") else None


def set_cell_borders(cell, top=None, bottom=None):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, spec in (("top", top), ("bottom", bottom)):
        if spec is None:
            continue
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(spec[0]))
        e.set(qn("w:color"), spec[1])
        borders.append(e)
    for edge in ("left", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tcPr.append(borders)


def box_borders(cell, edges=("top", "bottom", "left", "right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom", "left", "right"):
        e = OxmlElement(f"w:{edge}")
        if edge in edges:
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "6")
            e.set(qn("w:color"), "000000")
        else:
            e.set(qn("w:val"), "nil")
        borders.append(e)
    tcPr.append(borders)


def cell_text(cell, text, bold=False, italic=False, size=10.5, align=None, space=2):
    cell.text = ""
    par = cell.paragraphs[0]
    par.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.space_after = Pt(space)
    par.paragraph_format.space_before = Pt(space)
    r = par.add_run(text)
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = RGBColor(0, 0, 0)
    return par


def build_coversheet(doc):
    """The institute's Assignment Cover Sheet, reproduced as the first page."""
    t = plain(doc, "Assignment Cover Sheet", size=16, bold=True, italic=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    t.runs[0].font.name = "Times New Roman"

    logo = ROOT / "logos" / "learnkey.png"
    if logo.exists():
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.paragraph_format.space_after = Pt(12)
        lp.add_run().add_picture(str(logo), height=Cm(1.75))

    i = plain(doc, "This cover sheet must be completed and added to the front of every assignment",
              size=11, bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
    i.runs[0].font.name = "Times New Roman"

    fields = [
        ("Learner Name and Surname:", "Imran Hossain Chowdhury", False),
        ("Learner Registration No.", "11248", False),
        ("Study Centre Name", "Learn Key Institute", True),
        ("Qualification Title", "Undergraduate Diploma in Software Design MQF (Lv.5) — Group A", False),
        ("Unit Title", "DevOps Delivery Strategy — Individual Project & Portfolio", False),
        ("Submission Date", "29/08/2026", False),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    for label, value, bold_value in fields:
        r = tbl.add_row()
        cell_text(r.cells[0], label, bold=True, italic=True)
        cell_text(r.cells[1], value, bold=bold_value)
        box_borders(r.cells[0]); box_borders(r.cells[1])
    tbl.columns[0].width = Cm(6.2)
    tbl.columns[1].width = Cm(10.3)

    dec = doc.add_table(rows=1, cols=1).rows[0].cells[0]
    dec.text = ""
    head = dec.paragraphs[0]
    head.paragraph_format.space_after = Pt(4)
    hr = head.add_run("Declaration of authenticity:")
    hr.bold = True; hr.italic = True; hr.font.size = Pt(10.5); hr.font.name = "Times New Roman"
    for n, line in enumerate([
        "I declare that the attached submission is my own original work. No significant part of it has been "
        "submitted for any other assignment and I have acknowledged in my notes and bibliography all written "
        "and electronic sources used.",
        "I acknowledge that my assignment will be subject to electronic scrutiny for academic honesty.",
        "I understand that failure to meet these guidelines may instigate the centre's malpractice procedures "
        "and risk failure of the unit and/or qualification.",
    ], start=1):
        par = dec.add_paragraph()
        par.paragraph_format.space_after = Pt(4)
        par.paragraph_format.left_indent = Cm(0.6)
        r = par.add_run(f"{n}.  {line}")
        r.bold = True; r.italic = True; r.font.size = Pt(10.5); r.font.name = "Times New Roman"
    box_borders(dec)

    signature = next((ROOT / "logos" / n for n in ("signature.png", "signature.jpg", "signature.jpeg")
                      if (ROOT / "logos" / n).exists()), None)

    sig = doc.add_table(rows=1, cols=2).rows[0]
    for idx, label in enumerate(("Learner signature", "Tutor signature")):
        cell = sig.cells[idx]
        cell.text = ""
        spacer = cell.paragraphs[0]
        spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if idx == 0 and signature is not None:
            spacer.paragraph_format.space_after = Pt(0)
            spacer.add_run().add_picture(str(signature), height=Cm(1.2))
        else:
            spacer.paragraph_format.space_after = Pt(20)
        for text, is_rule in (("__________________________", True), (label, False), ("Date:", False)):
            par = cell.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.space_after = Pt(1)
            r = par.add_run(text)
            r.font.size = Pt(10.5)
            r.font.name = "Times New Roman"
            r.bold = not is_rule
            r.italic = not is_rule
        box_borders(cell)

    note = doc.add_table(rows=1, cols=1).rows[0].cells[0]
    cell_text(note, "Note: Assignments must be submitted in typed PDF format only; handwritten assignments "
                    "are not accepted.", size=10)
    box_borders(note)

    doc.add_page_break()


def add_runs(par, runs, size=Pt(10.5), colour=INK, italic=False):
    for text, fmt in runs:
        r = par.add_run(text)
        r.font.size = size
        r.font.color.rgb = colour
        r.font.name = MONO_FONT if "code" in fmt else BODY_FONT
        if "code" in fmt:
            r.font.size = Pt(size.pt - 1.2)
        r.bold = "b" in fmt
        r.italic = italic or ("i" in fmt)


def bookmark(par, name):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(abs(hash(name)) % 10000))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(abs(hash(name)) % 10000))
    par._p.insert(0, start)
    par._p.append(end)


def add_internal_link(par, anchor, runs_spec):
    link = OxmlElement("w:hyperlink")
    link.set(qn("w:anchor"), anchor)
    for text, bold, colour, size in runs_spec:
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), BODY_FONT)
        rf.set(qn("w:hAnsi"), BODY_FONT)
        rPr.append(rf)
        if bold:
            rPr.append(OxmlElement("w:b"))
        c = OxmlElement("w:color")
        c.set(qn("w:val"), colour)
        rPr.append(c)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        link.append(r)
    par._p.append(link)


def add_page_field(par):
    for instr in ("PAGE",):
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), instr)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "1"
        r.append(t)
        fld.append(r)
        par._p.append(fld)


def hrule(par, size=8, colour=DARK):
    pPr = par._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), "4")
    b.set(qn("w:color"), colour)
    borders.append(b)
    pPr.append(borders)


def left_bar(par, colour="7B2D26"):
    pPr = par._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    b = OxmlElement("w:left")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "12")
    b.set(qn("w:space"), "8")
    b.set(qn("w:color"), colour)
    borders.append(b)
    pPr.append(borders)


# --------------------------------------------------------------------------
# Document assembly
# --------------------------------------------------------------------------
PARTS = [
    ("part1", "Part One", "Current State and the Case for Change"),
    ("part2", "Part Two", "Version Control and Collaborative Workflow"),
    ("part3", "Part Three", "Continuous Delivery and Quality"),
    ("part4", "Part Four", "Environments, Infrastructure and Containers"),
    ("part5", "Part Five", "Observability, Security and Reliability"),
    ("part6", "Part Six", "Cloud, Cost and Leading the Change"),
]

TOC = [
    ("part1", "1", "Current State and the Case for Change",
     "Value-stream analysis, waste and handoffs, a critique of CALMS, and the baseline metrics."),
    ("part2", "2", "Version Control and Collaborative Workflow",
     "Branching strategy, review as knowledge transfer, release tagging and traceability."),
    ("part3", "3", "Continuous Delivery and Quality",
     "Pipeline design, the test portfolio, quality gates, automation priorities, release and rollback."),
    ("part4", "4", "Environments, Infrastructure and Containers",
     "Infrastructure as code, immutability, defence against drift, and the orchestration decision."),
    ("part5", "5", "Observability, Security and Reliability",
     "Telemetry, service level objectives, shifted-left security, incident response and recovery."),
    ("part6", "6", "Cloud, Cost and Leading the Change",
     "Cloud environment management, FinOps, the adoption roadmap and how success is measured."),
    ("appendix", "A", "Appendices",
     "A — supporting artefacts and how to verify them. B — references."),
]


def base_styles(doc):
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = Pt(10.5)
    st.font.color.rgb = INK
    pf = st.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.12
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        s = doc.styles[name]
        s.font.name = HEAD_FONT
        s.font.color.rgb = INK
        s.font.bold = True
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(11.5)
    doc.styles["Heading 3"].font.size = Pt(10.5)


def plain(doc, text="", size=10.5, colour=INK, bold=False, italic=False,
          align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = colour
        r.font.name = BODY_FONT
        r.bold = bold
        r.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    return p


def build_cover(doc):
    p = plain(doc, "LEARNKEY INSTITUTE  ·  MALTA CAMPUS", size=9, colour=MUTED,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    p = plain(doc, "Undergraduate Diploma in Software Design  ·  MQF Level 5  ·  Group A\nModule: Introduction to DevOps: Principles and Practices",
              size=9, colour=MUTED, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    hrule(p, size=12)

    for _ in range(7):
        plain(doc, space_after=0)

    plain(doc, "Scenario B — The Growing SaaS Start-up Outgrowing Its Scripts",
          size=12, colour=ACCENT, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
    t = plain(doc, "A DevOps Delivery Strategy for Nimbus Ltd.", size=26, bold=True,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_after=12)
    t.paragraph_format.line_spacing = 1.0
    plain(doc,
          "Nimbus is a fifteen-person start-up whose delivery process has not grown with it. Releases take "
          "seventeen days, one engineer holds all production knowledge, environments drift, and the cloud bill "
          "has no owner. This report analyses that delivery system, sets a measurement baseline, and designs an "
          "end-to-end strategy covering version control, continuous delivery, infrastructure as code, "
          "observability, security, cost and the twelve-month plan to adopt it.",
          size=10, colour=MUTED, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

    for _ in range(6):
        plain(doc, space_after=0)

    meta = doc.add_table(rows=0, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in (
        ("Author", "Imran Hossain Chowdhury"),
        ("Registration no.", "11248"),
        ("Unit", "DevOps Delivery Strategy — Individual Project & Portfolio"),
        ("Word count", f"{WORD_COUNT} words (excluding diagrams, code and appendices)"),
        ("Supporting repository", "github.com/imranneta5555/DevOps"),
        ("Date", "29 August 2026"),
    ):
        row = meta.add_row()
        for idx, text in enumerate((label, value)):
            cell = row.cells[idx]
            cell.text = ""
            par = cell.paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            par.paragraph_format.space_after = Pt(2)
            r = par.add_run(text)
            r.font.size = Pt(9.5)
            r.font.name = BODY_FONT
            r.font.color.rgb = MUTED if idx == 0 else INK
            set_cell_borders(cell)
    meta.columns[0].width = Cm(4.2)
    meta.columns[1].width = Cm(12.3)
    doc.add_page_break()


def build_toc(doc):
    h = plain(doc, "Contents", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
    hrule(h, size=12)
    for anchor, num, title, desc in TOC:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        add_internal_link(p, anchor, [(f"{num}   ", True, "7B2D26", 10.5),
                                      (title, True, "14161A", 10.5)])
        d = plain(doc, desc, size=9, colour=MUTED, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=7)
        d.paragraph_format.left_indent = Cm(0.9)
        hrule(d, size=4, colour=RULE)

    plain(doc, space_after=4)
    plain(doc,
          "Figures.  1 Current-state value stream · 2 Target delivery pipeline · 3 A pipeline run · "
          "4 Target environment architecture · 5 Error-budget policy and incident flow · "
          "6 The service the artefacts deploy",
          size=9, colour=MUTED, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3)
    plain(doc,
          "Tables.  1 Baseline metrics and targets · 2 Test portfolio · 3 Automation return on investment · "
          "4 Defence against drift · 5 Adoption roadmap",
          size=9, colour=MUTED, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0)
    doc.add_page_break()


def write_blocks(doc, blocks, anchor=None, label=None, title=None, lead=None):
    if title:
        lp = plain(doc, label.upper(), size=9, colour=ACCENT, bold=True,
                   align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
        bookmark(lp, anchor)
        h = doc.add_heading(title, level=1)
        h.paragraph_format.space_after = Pt(2)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lead_p = plain(doc, lead or "", size=9, colour=MUTED, italic=True,
                       align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
        hrule(lead_p, size=4, colour=RULE)

    for b in blocks:
        if b.kind == "h2":
            continue
        if b.kind == "lead":
            continue
        if b.kind == "h3":
            p = doc.add_heading("", level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, b.runs, size=Pt(11.5))
        elif b.kind == "h4":
            p = doc.add_heading("", level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs(p, b.runs, size=Pt(10.5))
        elif b.kind == "p":
            p = doc.add_paragraph()
            add_runs(p, b.runs)
        elif b.kind == "list":
            for item in b.items:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(3)
                add_runs(p, item)
        elif b.kind == "pre":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(b.text)
            r.font.name = MONO_FONT
            r.font.size = Pt(8)
            left_bar(p, "D6D9DD")
        elif b.kind == "note":
            for i, nb in enumerate(b.body):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                add_runs(p, nb.runs, size=Pt(9.8))
                left_bar(p)
        elif b.kind == "table":
            write_table(doc, b)
        elif b.kind == "figure":
            write_figure(doc, b)


def write_table(doc, b):
    if b.caption:
        plain(doc, b.caption, size=8.7, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_after=3, space_before=6)
    cols = len(b.head) if b.head else len(b.rows[0])
    table = doc.add_table(rows=0, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if b.head:
        row = table.add_row()
        for i, cell_spec in enumerate(b.head):
            cell = row.cells[i]
            cell.text = ""
            par = cell.paragraphs[0]
            par.paragraph_format.space_after = Pt(3)
            par.paragraph_format.space_before = Pt(3)
            par.alignment = WD_ALIGN_PARAGRAPH.RIGHT if cell_spec["num"] else WD_ALIGN_PARAGRAPH.LEFT
            r = par.add_run(cell_spec["text"])
            r.bold = True
            r.font.size = Pt(8.7)
            r.font.name = BODY_FONT
            set_cell_borders(cell, top=(10, DARK), bottom=(6, DARK))

    for ri, r_spec in enumerate(b.rows):
        row = table.add_row()
        last = ri == len(b.rows) - 1
        for i, cell_spec in enumerate(r_spec):
            cell = row.cells[i]
            cell.text = ""
            par = cell.paragraphs[0]
            par.paragraph_format.space_after = Pt(3)
            par.paragraph_format.space_before = Pt(3)
            par.alignment = WD_ALIGN_PARAGRAPH.RIGHT if cell_spec["num"] else WD_ALIGN_PARAGRAPH.LEFT
            r = par.add_run(cell_spec["text"])
            r.font.size = Pt(8.7)
            r.font.name = BODY_FONT
            set_cell_borders(cell, bottom=(10, DARK) if last else (4, RULE))
    plain(doc, space_after=4)


def write_figure(doc, b):
    key = None
    for name, (path, width) in FIGURE_IMAGES.items():
        if b.src and Path(b.src).stem == Path(path).stem:
            key = (path, width)
            break
    if key is None:
        # Inline SVG figure: match on the caption's figure number.
        text = "".join(t for t, _ in b.caption)
        m = re.match(r"Figure (\d)", text)
        order = {"1": "fig1-value-stream", "2": "fig2-pipeline", "4": "fig3-architecture",
                 "5": "fig4-error-budget"}
        if m and m.group(1) in order:
            key = FIGURE_IMAGES[order[m.group(1)]]
    if key is None:
        return
    path, width = key
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(ROOT / path), width=width)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_after = Pt(10)
    add_runs(cap, b.caption, size=Pt(8.7), colour=MUTED)


def add_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run("DevOps Delivery Strategy — Nimbus (Scenario B)\t")
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED
    r.font.name = BODY_FONT
    r2 = p.add_run("Page ")
    r2.font.size = Pt(8)
    r2.font.color.rgb = MUTED
    r2.font.name = BODY_FONT
    add_page_field(p)


FIGURE_RE = re.compile(r"<!--\s*FIGURE:([a-z0-9\-]+)\s*-->")


def inline_figures(html: str) -> str:
    """Same substitution assemble.py performs when building report.html."""

    def replace(match: re.Match) -> str:
        fragment = ROOT / "figures" / f"{match.group(1)}.html"
        return fragment.read_text(encoding="utf-8") if fragment.exists() else ""

    return FIGURE_RE.sub(replace, html)


def parse(path):
    p = Parser()
    p.feed(inline_figures(Path(path).read_text(encoding="utf-8")))
    return p.blocks


def main():
    doc = Document()
    base_styles(doc)
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.25)
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)

    build_coversheet(doc)
    build_cover(doc)
    build_toc(doc)

    for i, (anchor, label, title) in enumerate(PARTS):
        blocks = parse(ROOT / "parts" / f"{anchor}.html")
        lead = next((("".join(t for t, _ in b.runs)) for b in blocks if b.kind == "lead"), "")
        write_blocks(doc, blocks, anchor=anchor, label=label, title=title, lead=lead)
        doc.add_page_break()

    blocks = parse(ROOT / "parts" / "appendix.html")
    lead = next((("".join(t for t, _ in b.runs)) for b in blocks if b.kind == "lead"), "")
    write_blocks(doc, blocks, anchor="appendix", label="Appendix A",
                 title="Supporting Artefacts", lead=lead)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"built: {OUT.relative_to(ROOT)} ({OUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
